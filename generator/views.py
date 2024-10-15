from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import Lender, Borrower, DocumentTemplate, ExcelUpload
from openpyxl import load_workbook
from docx import Document
from django.contrib.admin.views.decorators import staff_member_required
from .forms import DocumentTemplateForm
from django.http import HttpResponse
from .forms import ExcelUploadForm, DocumentGenerationForm
import json
from django import forms
from io import BytesIO


class LenderBorrowerSelectionForm(forms.Form):
    lender = forms.ChoiceField(choices=[])
    borrower = forms.ChoiceField(choices=[])
    documents = forms.ModelMultipleChoiceField(queryset=DocumentTemplate.objects.all(), widget=forms.CheckboxSelectMultiple)

def select_lender_borrower(request):
    # Load the extracted data from session
    extracted_data = json.loads(request.session.get('extracted_data', '[]'))
    
    if not extracted_data:
        return HttpResponse("No data found, please upload an Excel file first.")

    # Get unique lender and borrower pairs
    lender_choices = [(i, data['lender']) for i, data in enumerate(extracted_data)]
    borrower_choices = [(i, data['borrower']) for i, data in enumerate(extracted_data)]

    if request.method == 'POST':
        form = LenderBorrowerSelectionForm(request.POST)
        form.fields['lender'].choices = lender_choices
        form.fields['borrower'].choices = borrower_choices
        
        if form.is_valid():
            selected_lender = int(form.cleaned_data['lender'])
            selected_borrower = int(form.cleaned_data['borrower'])
            documents = form.cleaned_data['documents']
            
            # Get selected row data
            data = extracted_data[selected_lender]

            # We'll store the output file in memory and send it as a response
            memory_file = BytesIO()

            # For demonstration, we assume one document is selected. Loop through if multiple documents.
            for doc_template in documents:
                template_path = doc_template.template_file.path
                doc = Document(template_path)

                # Replace placeholders with data
                for paragraph in doc.paragraphs:
                    for key, value in data.items():
                        if f'{{{{{key}}}}}' in paragraph.text:  # Match placeholder format {{key}}
                            paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))

                # Save the document to memory
                doc.save(memory_file)

            # Reset the file pointer to the beginning of the file
            memory_file.seek(0)

            # Prepare response for download
            response = HttpResponse(memory_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{doc_template.name}_{data["borrower"]}.docx"'

            return response

    else:
        form = LenderBorrowerSelectionForm()
        form.fields['lender'].choices = lender_choices
        form.fields['borrower'].choices = borrower_choices

    return render(request, 'generator/select_lender_borrower.html', {'form': form})


def upload_excel(request):
    if request.method == 'POST':
        form = ExcelUploadForm(request.POST, request.FILES)
        if form.is_valid():
            # Load the Excel file
            excel_file = request.FILES['excel_file']
            wb = load_workbook(excel_file)
            sheet = wb.active

            # Extract data from Excel and store it in session
            extracted_data = []
            for row in sheet.iter_rows(min_row=2, values_only=True):  # Skip the header
                row_data = {
                    'lender': row[0],
                    'borrower': row[1],
                    'principal': row[2],
                    'interest_rate': row[3],
                    'funding_date': row[4],
                    'maturity_date': row[5],
                    'property': row[6],
                    'guarantor': row[7],
                }
                extracted_data.append(row_data)
            
            # Store extracted data in session
            request.session['extracted_data'] = json.dumps(extracted_data)
            
            # Redirect to the selection page
            return redirect('select_lender_borrower')
    else:
        form = ExcelUploadForm()

    return render(request, 'generator/upload_excel.html', {'form': form})


@staff_member_required
def upload_template(request):
    if request.method == 'POST':
        form = DocumentTemplateForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('upload_and_generate')
    else:
        form = DocumentTemplateForm()

    return render(request, 'generator/upload_template.html', {'form': form})


def upload_and_generate(request):
    if request.method == 'POST':
        form = DocumentGenerationForm(request.POST, request.FILES)
        if form.is_valid():
            # Get the uploaded Excel file
            excel_file = form.cleaned_data['excel_file']
            documents = form.cleaned_data['documents']

            # Load the Excel file into memory
            wb = load_workbook(excel_file)
            sheet = wb.active

            # Iterate over each row in the Excel file starting from row 2 (assuming row 1 is headers)
            for row in sheet.iter_rows(min_row=2, values_only=True):
                data = {
                    'lender': row[0],
                    'borrower': row[1],
                    'principal': row[2],
                    'interest_rate': row[3],
                    'funding_date': row[4],
                    'maturity_date': row[5],
                    'property': row[6],
                    'guarantor': row[7],
                }

                # Generate documents for this row
                for doc_template in documents:
                    template_path = doc_template.template_file.path
                    doc = Document(template_path)

                    # Replace placeholders in the document with values from the current row
                    for paragraph in doc.paragraphs:
                        for key, value in data.items():
                            if f'{{{{{key}}}}}' in paragraph.text:  # Match placeholder format {{key}}
                                paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))

                    # Save the document with a unique filename based on borrower name or some unique identifier
                    output_path = f'generated_docs/{doc_template.name}_{data["borrower"]}.docx'
                    doc.save(output_path)

            return HttpResponse("Documents Generated Successfully")
    else:
        form = DocumentGenerationForm()

    return render(request, 'generator/upload_and_generate.html', {'form': form})